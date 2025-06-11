from django.shortcuts import render, get_object_or_404, redirect
from .models import Site, ElementControl, Local, Agent, GrilleControl, ControlGrille, Critere
from django.contrib.auth.decorators import login_required
from docx import Document
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from django.utils.timezone import datetime
from collections import defaultdict
from django.db.models import Sum, Count
from django.db.models.functions import ExtractMonth, ExtractYear
from django.http import JsonResponse



import matplotlib.pyplot as plt
import io
import urllib, base64
from django.db.models import Avg


# Create your views here.

@login_required
def acceuil(request) :
    user = request.user
    if user.role in ['admin', 'controleur'] :
         site = Site.objects.all()
    else : 
        site = user.sites_autorises.all()

    context = {
        'site' : site,
    }

    return render(request, "relotagrid/dash.html", context)

@login_required
def vue_app(request, slug) :
    site = get_object_or_404(Site, slug = slug)
    
    context = {
        "site" : site,
    }

    return render(request, 'relotagrid/vue_app.html', context)

@login_required
def detail_site(request, slug) :
    site = get_object_or_404(Site, slug = slug)
    locales = Local.objects.filter(site = site)
    context = {
        'site' : site,
        'locales' : locales,
    }

    return render(request, 'relotagrid/dash_localSite.html', context)

@login_required
def liste_agent(request) :
    agent = Agent.objects.all()
    context = {
        'agents' : agent,
    }

    return render(request, 'relotagrid/agent.html', context)

@login_required
def agent_detail(request, id):
    agent = get_object_or_404(Agent, id=id)

    # Ex. : données fictives pour le graphique
    graph_data = {
        'labels': ['Jan', 'Feb', 'Mar', 'Apr'],
        'values': [5, 8, 3, 7]
    }

    context = {
        'agent' : agent,
        'graph_data' : graph_data,
    }

    return render(request, 'relotagrid/detail_agent.html', context)

@login_required
def export_agents_docx(request):
    document = Document()
    document.add_heading("Fiche des Agents - Et leurs performances", 0)

    agents = Agent.objects.all()

    # Créer le tableau
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Style de l'entête
    hdr_cells = table.rows[0].cells
    headers = ["Nom", "Prénom", "Téléphone", "Performance de travail"]

    for i, text in enumerate(headers):
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)  # Texte blanc
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), "0D6EFD")  # vert Bootstrap
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

    # Lignes de données
    for agent in agents:
        row_cells = table.add_row().cells
        row_cells[0].text = agent.nom
        row_cells[1].text = agent.prenom
        row_cells[2].text = agent.telephone
        row_cells[3].text = ""  # Performance vide

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # Réponse HTTP avec le fichier
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="fiche_agents.docx"'
    document.save(response)
    return response

@login_required
def controller(request, id):
    local = Local.objects.get(id=id)
    site = local.site  # site du local
    agents = Agent.objects.filter(site__id=site.id)  # relation ManyToMany

    context = {
        'local': local,
        'agents': agents,
    }
    return render(request, 'relotagrid/controllerLocale.html', context)


@login_required
def creer_grille(request, slug) : 
    site = get_object_or_404(Site, slug = slug)
    locales = Local.objects.get(site = site)
    utilisateur = request.user
    grille = GrilleControl.objects.create(local = locales, utilisateur = utilisateur)

    print(grille.local.local_type)
    context = {
        'grille' : grille,
    }

    return render(request, 'relotagrid/dash_seeGrilleControl.html', context)


######## Creation de la grille et le controle #####################
@login_required
def enregistrer_grille(request, local_id):
    if request.method == "POST":
        local = get_object_or_404(Local, id=local_id)

        # Création de la grille
        grille = GrilleControl.objects.create(
            local=local,
            utilisateur=request.user
        )

        # Récupération de l'agent
        agent_id = request.POST.get("agent")
        if not agent_id:
            return HttpResponse("Aucun agent sélectionné", status=400)
        agent = get_object_or_404(Agent, id=agent_id)

        # Récupération de tous les éléments de contrôle
        controle_ids = request.POST.getlist("controle_ids[]")

        for controle_id in controle_ids:
            try:
                element = ElementControl.objects.get(id=controle_id)
            except ElementControl.DoesNotExist:
                continue

            # Parcours des critères de cet élément
            for critere in element.critere.all():
                note_key = f"note_{critere.id}"
                if note_key in request.POST:
                    try:
                        note = int(request.POST[note_key])
                    except ValueError:
                        continue  # Ignore si la note n'est pas valide

                    # Enregistrement dans ControlGrille
                    ControlGrille.objects.create(
                        grille=grille,
                        element_controle=element,
                        critere=critere,
                        note=note,
                        agent=agent
                    )

        return redirect('relotagrid:voir_controle', local.id)  # Redirection après enregistrement
    else:
        return HttpResponse("Méthode non autorisée", status=405)

 # Une fonction qui sert a visualiser les controles d'un local type donné  
@login_required  
def voir_controle(request, id) :
    locale = Local.objects.get(id=id)

    grille_list = GrilleControl.objects.prefetch_related('notes').filter(local = locale).order_by('-date','-heure')
    
    context = {
        'grille_list' : grille_list,
        'locale' : locale,
    }
    return render(request, "relotagrid/dash_allLocalControl.html", context)

# Fonction qui permet de visualiser la grille d'un controle 
@login_required
def voir_grille(request, id) :
    grille = GrilleControl.objects.get(id = id) 
    total_criteres = sum(ec.critere.count() for ec in grille.local.element_controle.all())
    objectif = total_criteres * 3
    context = {
        'grille' : grille,
        'objectif' : objectif,
    }
    return render(request, 'relotagrid/grille.html', context)

# Fonction qui va generer les graphe des grilles
@login_required
def performance_mensuelle_site(request, slug):
    site = get_object_or_404(Site, slug=slug)

    # Nouveau champ "mois_annee" sous format YYYY-MM
    mois_annee = request.GET.get('mois_annee')
    if mois_annee:
        try:
            date_obj = datetime.strptime(mois_annee, "%Y-%m")
            mois = date_obj.month
            annee = date_obj.year
        except ValueError:
            mois = datetime.now().month
            annee = datetime.now().year
    else:
        mois = datetime.now().month
        annee = datetime.now().year

    # Le reste inchangé
    agents = Agent.objects.filter(site=site)
    performances = {}

    for agent in agents:
        notes = ControlGrille.objects.filter(
            agent=agent,
            grille__date__month=mois,
            grille__date__year=annee,
            grille__local__site=site
        )

        total_notes = notes.aggregate(total=Sum('note'))['total'] or 0
        total_max = notes.aggregate(nombre=Count('id'))['nombre'] * 3

        pourcentage = round((total_notes / total_max) * 100, 2) if total_max > 0 else 0
        performances[f"{agent.nom} {agent.prenom}"] = pourcentage

    background_colors = [
        'rgba(75, 192, 192, 0.7)' if val >= 70 else 'rgba(255, 99, 132, 0.7)'
        for val in performances.values()
    ]

    context = {
        'site': site,
        'labels': list(performances.keys()),
        'values': list(performances.values()),
        'background_colors': background_colors,
        'mois': mois,
        'annee': annee,
    }

    return render(request, 'relotagrid/graphe_grille.html', context)
